"""
ATS-optimized resume DOCX renderer.

Renders structured resume JSON (from LLM rewrite) into a clean,
single-column, ATS-compatible DOCX using python-docx.

Template: ats_classic — maximally ATS-safe:
- Single column layout
- Standard section headings
- No tables, graphics, or complex formatting
- Clean sans-serif font (Calibri)
- Consistent spacing
"""
import io

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# ── Colours ──────────────────────────────────────────────────────────────

COLOR_PRIMARY = RGBColor(0x1a, 0x1a, 0x2e)      # Dark navy
COLOR_SECONDARY = RGBColor(0x16, 0x21, 0x3e)     # Slightly lighter
COLOR_BODY = RGBColor(0x2c, 0x3e, 0x50)          # Body text
COLOR_MUTED = RGBColor(0x7f, 0x8c, 0x8d)         # Dates, locations
COLOR_DIVIDER = RGBColor(0xbd, 0xc3, 0xc7)       # Section dividers


# ── Style setup ──────────────────────────────────────────────────────────

def _setup_styles(doc):
    """Configure document-level styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(13)

    # Section heading style
    if 'SectionHeading' not in [s.name for s in doc.styles]:
        heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.font.color.rgb = COLOR_PRIMARY
        heading_style.paragraph_format.space_before = Pt(10)
        heading_style.paragraph_format.space_after = Pt(4)
        heading_style.paragraph_format.line_spacing = Pt(15)


def _add_divider(doc):
    """Add a thin horizontal line after section headings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    # Use a border-bottom on the paragraph
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BDC3C7')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Section builders ─────────────────────────────────────────────────────

def _build_contact(doc, content):
    """Add contact header: name centered, details below."""
    contact = content.get('contact', {})
    name = contact.get('name', 'Candidate')

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(2)

    # Contact details line
    parts = []
    if contact.get('email'):
        parts.append(contact['email'])
    if contact.get('phone'):
        parts.append(contact['phone'])
    if contact.get('location'):
        parts.append(contact['location'])
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' | '.join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        p.paragraph_format.space_after = Pt(2)

    # Links line
    links = []
    if contact.get('linkedin'):
        links.append(contact['linkedin'])
    if contact.get('portfolio'):
        links.append(contact['portfolio'])
    if links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' | '.join(links))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        p.paragraph_format.space_after = Pt(4)

    _add_divider(doc)


def _build_summary(doc, content):
    """Add professional summary."""
    summary = content.get('summary', '')
    if not summary:
        return

    doc.add_paragraph('PROFESSIONAL SUMMARY', style='SectionHeading')
    _add_divider(doc)

    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(4)


def _build_experience(doc, content):
    """Add work experience section."""
    experience = content.get('experience', [])
    if not experience:
        return

    doc.add_paragraph('WORK EXPERIENCE', style='SectionHeading')
    _add_divider(doc)

    for exp in experience:
        # Job title
        p = doc.add_paragraph()
        run = p.add_run(exp.get('title', ''))
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_SECONDARY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

        # Company and location
        company_line = exp.get('company', '')
        if exp.get('location'):
            company_line += f' — {exp["location"]}'
        p = doc.add_paragraph()
        run = p.add_run(company_line)
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY
        p.paragraph_format.space_after = Pt(1)

        # Dates
        start = exp.get('start_date', '')
        end = exp.get('end_date', '')
        if start or end:
            date_line = f'{start} – {end}' if start and end else (start or end)
            p = doc.add_paragraph()
            run = p.add_run(date_line)
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.space_after = Pt(3)

        # Bullets
        for bullet in exp.get('bullets', []):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(bullet)
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BODY
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)


def _build_education(doc, content):
    """Add education section."""
    education = content.get('education', [])
    if not education:
        return

    doc.add_paragraph('EDUCATION', style='SectionHeading')
    _add_divider(doc)

    for edu in education:
        # Degree
        p = doc.add_paragraph()
        run = p.add_run(edu.get('degree', ''))
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_SECONDARY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        # Institution
        inst_line = edu.get('institution', '')
        if edu.get('location'):
            inst_line += f' — {edu["location"]}'
        p = doc.add_paragraph()
        run = p.add_run(inst_line)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY
        p.paragraph_format.space_after = Pt(1)

        # Year and GPA
        details = []
        if edu.get('year'):
            details.append(edu['year'])
        if edu.get('gpa'):
            details.append(f'GPA: {edu["gpa"]}')
        if details:
            p = doc.add_paragraph()
            run = p.add_run(' | '.join(details))
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.space_after = Pt(3)


def _build_skills(doc, content):
    """Add skills section — grouped by category."""
    skills = content.get('skills', {})
    if not skills:
        return

    has_any = any(skills.get(cat) for cat in ('technical', 'tools', 'soft'))
    if not has_any:
        return

    doc.add_paragraph('SKILLS', style='SectionHeading')
    _add_divider(doc)

    category_labels = {
        'technical': 'Technical Skills',
        'tools': 'Tools & Platforms',
        'soft': 'Soft Skills',
    }

    for cat_key, cat_label in category_labels.items():
        items = skills.get(cat_key, [])
        if items:
            p = doc.add_paragraph()
            run = p.add_run(f'{cat_label}: ')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_SECONDARY
            run = p.add_run(', '.join(items))
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BODY
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.15)


def _build_certifications(doc, content):
    """Add certifications section."""
    certs = content.get('certifications', [])
    if not certs:
        return

    doc.add_paragraph('CERTIFICATIONS', style='SectionHeading')
    _add_divider(doc)

    for cert in certs:
        parts = [cert.get('name', '')]
        if cert.get('issuer'):
            parts.append(f'— {cert["issuer"]}')
        if cert.get('year'):
            parts.append(f'({cert["year"]})')

        p = doc.add_paragraph()
        run = p.add_run(' '.join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY
        p.paragraph_format.space_after = Pt(2)


def _build_projects(doc, content):
    """Add projects section."""
    projects = content.get('projects', [])
    if not projects:
        return

    doc.add_paragraph('PROJECTS', style='SectionHeading')
    _add_divider(doc)

    for proj in projects:
        # Project name
        p = doc.add_paragraph()
        run = p.add_run(proj.get('name', ''))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_SECONDARY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        # Description
        if proj.get('description'):
            p = doc.add_paragraph()
            run = p.add_run(proj['description'])
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BODY
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)

        # Technologies
        techs = proj.get('technologies', [])
        if techs:
            p = doc.add_paragraph()
            run = p.add_run(f'Technologies: {", ".join(techs)}')
            run.font.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)

        # URL
        if proj.get('url'):
            p = doc.add_paragraph()
            run = p.add_run(proj['url'])
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(2)


# ── Main renderer ────────────────────────────────────────────────────────

def render_resume_docx(resume_content: dict) -> bytes:
    """
    Render structured resume JSON into an ATS-optimized DOCX.

    Args:
        resume_content: Validated resume JSON from LLM rewrite.

    Returns:
        DOCX bytes.
    """
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    _setup_styles(doc)

    # Build all sections in order
    _build_contact(doc, resume_content)
    _build_summary(doc, resume_content)
    _build_experience(doc, resume_content)
    _build_education(doc, resume_content)
    _build_skills(doc, resume_content)
    _build_certifications(doc, resume_content)
    _build_projects(doc, resume_content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
