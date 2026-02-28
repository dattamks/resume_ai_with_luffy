"""
Executive resume DOCX renderer.

Formal, authoritative layout with dark charcoal header, serif headings,
and a professional tone suited for senior/C-level positions.

Template: executive — formal, single-column, serif accents, dark header.
"""
import io

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# ── Colours ──────────────────────────────────────────────────────────────

COLOR_PRIMARY = RGBColor(0x1b, 0x26, 0x31)      # Dark charcoal
COLOR_SECONDARY = RGBColor(0x2c, 0x3e, 0x50)    # Slate
COLOR_BODY = RGBColor(0x2c, 0x3e, 0x50)          # Body
COLOR_MUTED = RGBColor(0x7f, 0x8c, 0x8d)         # Muted
COLOR_DIVIDER = RGBColor(0x1b, 0x26, 0x31)       # Dark divider


def _safe(text) -> str:
    if not text:
        return ''
    s = str(text)
    s = ''.join(c for c in s if c in ('\t', '\n', '\r') or (ord(c) >= 32))
    return s


def _setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = Pt(14)

    try:
        doc.styles['SectionHeading']
    except KeyError:
        heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.font.color.rgb = COLOR_PRIMARY
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(2)
        heading_style.paragraph_format.line_spacing = Pt(15)


def _add_divider(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B2631')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_contact(doc, content):
    contact = content.get('contact', {})
    name = _safe(contact.get('name', 'Candidate')).upper()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(4)

    parts = []
    if contact.get('email'):
        parts.append(_safe(contact['email']))
    if contact.get('phone'):
        parts.append(_safe(contact['phone']))
    if contact.get('location'):
        parts.append(_safe(contact['location']))
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' | '.join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        p.paragraph_format.space_after = Pt(2)

    links = []
    if contact.get('linkedin'):
        links.append(_safe(contact['linkedin']))
    if contact.get('portfolio'):
        links.append(_safe(contact['portfolio']))
    if links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' | '.join(links))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        p.paragraph_format.space_after = Pt(4)

    _add_divider(doc)


def _build_summary(doc, content):
    summary = _safe(content.get('summary', ''))
    if not summary:
        return
    doc.add_paragraph('EXECUTIVE SUMMARY', style='SectionHeading')
    _add_divider(doc)
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'


def _build_experience(doc, content):
    experience = content.get('experience', [])
    if not experience:
        return
    doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='SectionHeading')
    _add_divider(doc)
    for exp in experience:
        p = doc.add_paragraph()
        run = p.add_run(_safe(exp.get('title', '')))
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_PRIMARY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

        company_line = _safe(exp.get('company', ''))
        if exp.get('location'):
            company_line += f' — {_safe(exp["location"])}'
        p = doc.add_paragraph()
        run = p.add_run(company_line)
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_BODY
        p.paragraph_format.space_after = Pt(1)

        start = _safe(exp.get('start_date', ''))
        end = _safe(exp.get('end_date', ''))
        if start or end:
            date_line = f'{start} – {end}' if start and end else (start or end)
            p = doc.add_paragraph()
            run = p.add_run(date_line)
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.space_after = Pt(3)

        for bullet in exp.get('bullets', []):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(_safe(bullet))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_BODY
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)


def _build_education(doc, content):
    education = content.get('education', [])
    if not education:
        return
    doc.add_paragraph('EDUCATION', style='SectionHeading')
    _add_divider(doc)
    for edu in education:
        p = doc.add_paragraph()
        run = p.add_run(_safe(edu.get('degree', '')))
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_PRIMARY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        inst_line = _safe(edu.get('institution', ''))
        if edu.get('location'):
            inst_line += f' — {_safe(edu["location"])}'
        p = doc.add_paragraph()
        run = p.add_run(inst_line)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_BODY
        p.paragraph_format.space_after = Pt(1)

        details = []
        if edu.get('year'):
            details.append(_safe(edu['year']))
        if edu.get('gpa'):
            details.append(f'GPA: {_safe(edu["gpa"])}')
        if details:
            p = doc.add_paragraph()
            run = p.add_run(' | '.join(details))
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.space_after = Pt(3)


def _build_skills(doc, content):
    skills = content.get('skills', {})
    if not skills:
        return
    has_any = any(skills.get(cat) for cat in ('technical', 'tools', 'soft'))
    if not has_any:
        return
    doc.add_paragraph('CORE COMPETENCIES', style='SectionHeading')
    _add_divider(doc)
    category_labels = {'technical': 'Technical', 'tools': 'Tools & Platforms', 'soft': 'Leadership & Soft Skills'}
    for cat_key, cat_label in category_labels.items():
        items = skills.get(cat_key, [])
        if items:
            p = doc.add_paragraph()
            run = p.add_run(f'{cat_label}: ')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_PRIMARY
            run = p.add_run(', '.join(_safe(s) for s in items))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_BODY
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.15)


def _build_certifications(doc, content):
    certs = content.get('certifications', [])
    if not certs:
        return
    doc.add_paragraph('CERTIFICATIONS & CREDENTIALS', style='SectionHeading')
    _add_divider(doc)
    for cert in certs:
        parts = [_safe(cert.get('name', ''))]
        if cert.get('issuer'):
            parts.append(f'— {_safe(cert["issuer"])}')
        if cert.get('year'):
            parts.append(f'({_safe(cert["year"])})')
        p = doc.add_paragraph()
        run = p.add_run(' '.join(parts))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_BODY
        p.paragraph_format.space_after = Pt(2)


def _build_projects(doc, content):
    projects = content.get('projects', [])
    if not projects:
        return
    doc.add_paragraph('KEY PROJECTS', style='SectionHeading')
    _add_divider(doc)
    for proj in projects:
        p = doc.add_paragraph()
        run = p.add_run(_safe(proj.get('name', '')))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_PRIMARY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        if proj.get('description'):
            p = doc.add_paragraph()
            run = p.add_run(_safe(proj['description']))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_BODY
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)

        techs = proj.get('technologies', [])
        if techs:
            p = doc.add_paragraph()
            run = p.add_run(f'Technologies: {", ".join(_safe(t) for t in techs)}')
            run.font.italic = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)

        if proj.get('url'):
            p = doc.add_paragraph()
            run = p.add_run(_safe(proj['url']))
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = COLOR_MUTED
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(2)


# ── Main renderer ────────────────────────────────────────────────────────

def render_executive_docx(resume_content: dict) -> bytes:
    """Render structured resume JSON into an executive-styled DOCX."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    _setup_styles(doc)
    _build_contact(doc, resume_content)
    _build_summary(doc, resume_content)
    _build_experience(doc, resume_content)
    _build_education(doc, resume_content)
    _build_skills(doc, resume_content)
    _build_certifications(doc, resume_content)
    _build_projects(doc, resume_content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
